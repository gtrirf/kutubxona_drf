from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from django.conf import settings
import os
from io import BytesIO

def create_book_template(book):
    document = Document()

    document.add_heading(f'Book Details: {book.title}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    if book.image:
        image_path = os.path.join(settings.MEDIA_ROOT, book.image.name)
        if os.path.exists(image_path):
            document.add_picture(image_path, width=Inches(2.0))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=0, cols=2)

    def add_row(label, value):
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value if value else 'N/A'

    add_row('Title', book.title)
    add_row('Yozuvi', book.yozuvi.name_yozuv if book.yozuvi else "")
    add_row('Tili', book.tili.til if book.tili else "")
    add_row('Pages', str(book.pages))
    add_row('Nashriyot', book.nashriyot)
    add_row('Description', book.description)
    add_row('ISBN', book.isbn)
    add_row('Author', f'{book.author.first_name} {book.author.last_name}' if book.author else "")
    add_row('Realize Date', book.realize_date.strftime("%Y-%m-%d") if book.realize_date else "")
    add_row('Category', book.category.name if book.category else "")
    add_row('Quantity', str(book.quantity))
    add_row('Created At', book.created_at.strftime("%Y-%m-%d %H:%M:%S"))
    add_row('Updated At', book.updated_at.strftime("%Y-%m-%d %H:%M:%S"))

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream
